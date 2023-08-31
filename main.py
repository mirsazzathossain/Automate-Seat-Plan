"""
This script creates the exam seating documents for a given course.

Usage:
    python main.py --config config.yaml
"""

__author__ = "Mir Sazzat Hossain"

import os
import argparse
import yaml
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx2pdf import convert


def shuffle_students(students: pd.DataFrame) -> pd.DataFrame:
    """
    Shuffle the students in a dataframe.

    :param students: dataframe of students
    :type students: pd.DataFrame

    :return: shuffled dataframe of students
    :rtype: pd.DataFrame
    """
    for i in range(10):
        students = students.sample(frac=1).reset_index(drop=True)

    return students


def create_doc(
        students: pd.DataFrame,
        rooms: pd.DataFrame,
        course_code: str,
        exam_type: str,
        semester: str,
        year: int
) -> None:
    """
    Create the exam seating documents.

    :param students: dataframe of students
    :type students: pd.DataFrame
    :param rooms: dataframe of rooms
    :type rooms: pd.DataFrame
    :param course_code: course code
    :type course_code: str
    :param exam_type: exam type
    :type exam_type: str
    :param semester: semester
    :type semester: str
    :param year: year
    :type year: int
    """
    students_per_room = int(students.shape[0] // rooms.shape[0])

    # Create two documents
    doc_seat_plan = Document()
    doc_sign_sheet = Document()

    # Get the highest length of the student names
    max_name_length = students['Name'].apply(lambda x: len(x)).max()

    for doc in [doc_seat_plan, doc_sign_sheet]:
        # Set page size to A4 and margins to 1 inch
        section = doc.sections[0]
        section.page_height = Inches(11.69)
        section.page_width = Inches(8.27)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

        # Headder text
        text = f"{course_code} {exam_type} "
        if doc == doc_seat_plan:
            text += "Combined Seat Plan "
        else:
            text += "Signature Sheet "
        text += f"{semester} {year}"

        # Add the header
        header = section.header
        header_text = header.paragraphs[0]
        header_text.text = text

        # Set header font
        header_font = header_text.style.font
        header_font.name = 'Times New Roman'
        header_font.size = Pt(16)
        header_font.bold = True

        # Align header to center
        header_text.alignment = 1

        # Count extra students
        extra_students = students.shape[0] % rooms.shape[0]
        added_extra_students = 0

        # Add the students to the document
        for i in range(rooms['Rooms'].shape[0]):
            # Add the room number
            room_number = rooms['Rooms'][i]
            room_para = doc.add_paragraph()
            room_run = room_para.add_run(f"Room: {room_number}")
            room_run.bold = True
            room_run.font.size = Pt(14)
            room_run.font.name = 'Times New Roman'

            # Create the table
            table = doc.add_table(
                rows=students_per_room+1,
                cols=3 if doc == doc_seat_plan else 4
            )
            table.style = 'Table Grid'
            table.rows[0].cells[0].text = 'ID'
            table.rows[0].cells[1].text = 'Name'
            table.rows[0].cells[2].text = 'Section'
            if doc == doc_sign_sheet:
                table.rows[0].cells[3].text = 'Signature'

            # Add extra student if needed and update the start and end index
            start_index = i * students_per_room + added_extra_students
            if extra_students > 0:
                table.add_row()
                extra_students -= 1
                added_extra_students += 1
            end_index = (i+1) * students_per_room + added_extra_students

            # Get the students for the room
            room_students = students.iloc[start_index:end_index]

            # Sort the students by Section then ID
            room_students = room_students.sort_values(
                by=['Section', 'ID']
            ).reset_index(drop=True)

            # Add the students to the table
            for j in range(room_students.shape[0]):
                table.rows[j+1].cells[0].text = str(
                    room_students['ID'].iloc[j]
                )
                table.rows[j+1].cells[1].text = str(
                    room_students['Name'].iloc[j]
                )
                table.rows[j+1].cells[2].text = str(
                    room_students['Section'].iloc[j]
                )

                # Add the signature cell if needed
                if doc == doc_sign_sheet:
                    table.rows[j+1].cells[3].text = ''

            # Set the font of the table
            for row_index, row in enumerate(table.rows):
                for cell_index, cell in enumerate(row.cells):
                    if cell_index == 1:
                        cell.width = Cm(max_name_length * 0.3)
                    if doc == doc_sign_sheet:
                        if cell_index == 0 or cell_index == 2:
                            cell.width = Cm(1.2)
                        if cell_index == 3:
                            cell.width = Cm(max_name_length * 0.25)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.name = 'Times New Roman'
                            font.size = Pt(12)
                            font.bold = True if row_index == 0 else False
                        if cell_index != 1:
                            paragraph.alignment = 1

            # Align the table to center
            table.alignment = 1

            # Add a page break if needed
            if i != rooms.shape[0] - 1:
                doc.add_page_break()

    # Check if results folder exists
    if not os.path.exists('results'):
        os.makedirs('results')

    # Format the file names
    sign_sheet_file_name = f"{course_code}_{exam_type}_Signature_Sheet_"
    sign_sheet_file_name += f"{semester}_{year}.docx"
    seat_plan_file_name = f"{course_code}_{exam_type}_Combined_Seat_Plan_"
    seat_plan_file_name += f"{semester}_{year}.docx"

    # Save the documents
    doc_seat_plan.save(f"results/{seat_plan_file_name}")
    doc_sign_sheet.save(f"results/{sign_sheet_file_name}")

    # Convert the documents to pdf
    convert(f"results/{seat_plan_file_name}")
    convert(f"results/{sign_sheet_file_name}")


if __name__ == '__main__':
    parser = argparse.ArgumentParser()
    parser.add_argument('--config', type=str, default='config.yaml')
    args = parser.parse_args()

    with open(args.config, 'r') as f:
        config = yaml.load(f, Loader=yaml.FullLoader)

    # Read in the data
    rooms = pd.read_csv(config['rooms_file_path'])
    students = pd.read_csv(config['students_file_path'])

    # Get the config variables
    course_code = config['course_code']
    exam_type = config['exam_type']
    semester = config['semester']
    year = config['year']

    # Shuffle the students
    students = shuffle_students(students)

    # Create the documents
    create_doc(
        students,
        rooms,
        course_code,
        exam_type,
        semester,
        int(year)
    )
